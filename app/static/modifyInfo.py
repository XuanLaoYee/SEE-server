#修改某人数据（通过工号）
from docx import Document
import sys
title=sys.argv[1]
num=sys.argv[2]
try:
    Info=eval(sys.argv[3])
except Exception as ex:
    print(ex)
document = Document(title+'.docx')
tables=document.tables
table=tables[0]
rows = len(table.rows) 
#查找工号为num的人
for i in range(rows-1):
    row=table.rows[i+1]
    if(row.cells[3].text==str(num)):#找到
     #替换
        for j in range(len(table.columns)):
            row.cells[j].text=Info[j]
document.save(title+'.docx')

def sortInfo(title):
    document = Document(title+'.docx')
    paragraphs=document.paragraphs
    para=paragraphs[1]
    #得到了order：ABC
    tables=document.tables
    table=tables[0]#得到table
    order=(para.text)[-3:]#截取后三位
    rows=len(table.rows)
    cols=len(table.columns)#rows and cols 第一行是表头
    tmptable=[]#展示保存表格数据
    for i in range(rows-1):
        row=table.rows[i+1]
        tmp=[]
        for j in range(cols):
            tmp.append(row.cells[j].text)
        tmptable.append(tmp)     
    #得到所有的表格数据
    #TODO scan
    tmptable2=[]
    
    for i in range(len(order)):
        for j in range(rows-1):
            if(order[i]==tmptable[j][0]):
                tmptable2.append(tmptable[j])
    #TODO 装进去
    for i in range(rows-1):
        row=table.rows[i+1]
        for j in range(cols):
            row.cells[j].text=tmptable2[i][j]
    document.save(title+'.docx') 
sortInfo(title)