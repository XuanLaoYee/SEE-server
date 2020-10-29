#建立项目文件
from docx import Document 
import sys  
title=sys.argv[1]
order=sys.argv[2]
admin=sys.argv[3]
document=Document()
document.add_heading("project name:"+title, level=1)#添加标题
paragraph = document.add_paragraph("order:"+order)
paragraph=document.add_paragraph("admin:"+admin)
table = document.add_table(1, 5)
table.style='LightShading-Accent1' 
# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = '种类'
heading_cells[1].text = '子任务号'
heading_cells[2].text = '员工姓名'
heading_cells[3].text = '员工账号'
heading_cells[4].text = '完成时间'
document.save(title+'.docx')