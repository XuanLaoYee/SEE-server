#添加一条记录
from docx import Document
import sys
title=sys.argv[1]
try:
    Info=eval(sys.argv[2])
except Exception as ex:
    print(ex)
document = Document(title+'.docx')
tables=document.tables
table=tables[0]
cells = table.add_row().cells
cells[0].text = str(Info[0])
cells[1].text = Info[1]
cells[2].text = Info[2]
cells[3].text = Info[3]
cells[4].text = Info[4]
document.save(title+'.docx')

def sortInfo(title):
    document = Document(title+'.docx')
    paragraphs=document.paragraphs
    para=paragraphs[1]
    #得到了order：ABC
    tables=document.tables
    table=tables[0]#得到table
    order=(para.text)[-3:]#截取后三位
    rows=len(table.rows)
    cols=len(table.columns)#rows and cols 第一行是表头
    tmptable=[]#展示保存表格数据
    for i in range(rows-1):
        row=table.rows[i+1]
        tmp=[]
        for j in range(cols):
            tmp.append(row.cells[j].text)
        tmptable.append(tmp)     
    #得到所有的表格数据
    #TODO scan
    tmptable2=[]
    
    for i in range(len(order)):
        for j in range(rows-1):
            if(order[i]==tmptable[j][0]):
                tmptable2.append(tmptable[j])
    #TODO 装进去
    for i in range(rows-1):
        row=table.rows[i+1]
        for j in range(cols):
            row.cells[j].text=tmptable2[i][j]
    document.save(title+'.docx')
sortInfo(title)    